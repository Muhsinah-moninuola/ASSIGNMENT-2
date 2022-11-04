import docx
file = docx.Document('Nafsts.docx.doc')
for i in range(0,len(file.paragraphs),2):
    line=file.paragraphs[i].runs[0]
    line.text=line.text.replace('%20', ' ').replace('360p_','').replace(' .', '.')

file.save('Nafsts.docx')